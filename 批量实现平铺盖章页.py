# -*- coding: utf-8 -*-
"""
@Time ： 2022/11/21 19:40
@Auth ： 唐成
@File ：批量实现平铺盖章页.py
@IDE ：PyCharm

"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm
from add_float_picture import add_float_picture
from win32com import client
from time import sleep


'''需要修改签章页和说明性文件所在文件夹位置，文档尽量保存为docx'''
path_picture=r"C:\Users\唐成\Desktop\盖章页"
path_word=r"C:\Users\唐成\Desktop\说明性文件-阳投2\说明性文件-阳投2"
pics = os.listdir(path_picture)
words = os.listdir(path_word)
tc=len(pics)

nn=len(words)
for j in range(nn):
    print(words[j])


# 转换doc为pdf
def doc2pdf(fn):
    sleep(0.5)
    word = client.Dispatch("Word.Application")  # 打开word应用程序
    sleep(0.5)
    # for file in files:
    doc = word.Documents.Open(fn)  # 打开word文件
    sleep(0.5)
    doc.SaveAs("{}.pdf".format(fn[:-4]), 17)  # 另存为后缀为".pdf"的文件，其中参数17表示为pdf
    sleep(0.5)
    doc.Close()  # 关闭原来word文件
    word.Quit()


# 转换docx为pdf
def docx2pdf(fn):
    sleep(0.5)
    word = client.Dispatch("Word.Application")  # 打开word应用程序
    sleep(0.5)
    # for file in files:
    doc = word.Documents.Open(fn)  # 打开word文件
    sleep(0.5)
    doc.SaveAs("{}.pdf".format(fn[:-5]), 17)  # 另存为后缀为".pdf"的文件，其中参数17表示为pdf
    sleep(0.5)
    doc.Close()  # 关闭原来word文件
    word.Quit()


# # 删除word中所有签章页，一般不用
# for i in range(tc):
#     doc_save = path_word+"\\"+words[i]
#     doc = Document(doc_save)
#     for p in doc.paragraphs:
#         images = p._element.xpath('.//pic:pic')  # Get all pictures
#         for image in images:
#             # 在这个地方将段落里的内容保存下来 重新赋值 图片自己就会消失了
#             p.text = p.text
#             break
#     doc.save(doc_save)


for i in range(tc):

    if __name__ == '__main__':
        document = Document(path_word+"\\"+words[i])

        # add a floating picture
        p = document.paragraphs[-1] ##在最后一段插入
        add_float_picture(p, path_picture+"\\"+pics[i], width=Cm(21), pos_x=Cm(0), pos_y=Cm(0))

        # add text
        document.save(path_word+"\\"+words[i])

        #转换pdf
        fn = path_word + "\\" + words[i]
        docx2pdf(fn)











